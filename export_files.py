from tkinter import filedialog, messagebox
from docx import Document
import tkinter as tk

class ExportingFiles:
# ----------------------------------------------------------------------------------------------------
# Эта функция экспортирует полученную табличку в формате CSV-файла с предоставлением привилегии выбрать куда его положить.
    def export_results(self, notebook):
        # спрашиваем куда сохранять файл
        file_path = filedialog.asksaveasfilename(title="Save file", defaultextension=".csv",
                                                 filetypes=(("CSV files", "*.csv"), ("all files", "*.*")))
        if not file_path:
            return
        # конвертируем Дата фрейм в CSV-шник с кодировкой utf-8, дабы избежать корявых символов
        if notebook == "Station":
            self.new_df.to_csv(file_path, index=False, sep=';',encoding='utf-8')
            messagebox.showinfo("Экспорт файлов", "Экспорт завершен успешно")  
        elif notebook == "Horizon":
            self.horizon_table.to_csv(file_path, index=False, sep=';',encoding='utf-8')
            messagebox.showinfo("Экспорт файлов", "Экспорт завершен успешно")  
        elif notebook == "Chem elem":
            self.chem_elem_table.to_csv(file_path, index=False, sep=';',encoding='utf-8')
            messagebox.showinfo("Экспорт файлов", "Экспорт завершен успешно")  
         
# ----------------------------------------------------------------------------------------------------
    def export_report(self, notebook):
        if notebook == "Station":
            text = self.file_report_text.get("1.0", tk.END)
            doc = Document()
            doc.add_paragraph(text)

            # спрашиваем куда сохранять файл
            file_path = filedialog.asksaveasfilename(title="Save file", defaultextension=".docx",
                                                    filetypes=(("WORD files", "*.docx"),
                                                                ("DOC files", "*.doc"),
                                                                ("all files", "*.*")))
            if not file_path:
                return
            if file_path:
                doc.save(file_path)       
            messagebox.showinfo("Экспорт репорта", "Экспорт завершен успешно")

        elif notebook == "Horizon":
            text = self.file_report_text_horizon.get("1.0", tk.END)
            doc = Document()
            doc.add_paragraph(text)

            # спрашиваем куда сохранять файл
            file_path = filedialog.asksaveasfilename(title="Save file", defaultextension=".docx",
                                                    filetypes=(("WORD files", "*.docx"),
                                                                ("DOC files", "*.doc"),
                                                                ("all files", "*.*")))
            if not file_path:
                return
            if file_path:
                doc.save(file_path)       
            messagebox.showinfo("Экспорт репорта", "Экспорт завершен успешно")        

        elif notebook == "Chem elem":
            text = self.file_report_text_chem_elem.get("1.0", tk.END)
            doc = Document()
            doc.add_paragraph(text)

            # спрашиваем куда сохранять файл
            file_path = filedialog.asksaveasfilename(title="Save file", defaultextension=".docx",
                                                    filetypes=(("WORD files", "*.docx"),
                                                               ("DOC files", "*.doc"),
                                                               ("all files", "*.*")))
            if not file_path:
                return
            if file_path:
                doc.save(file_path)       
            messagebox.showinfo("Экспорт репорта", "Экспорт завершен успешно")     
